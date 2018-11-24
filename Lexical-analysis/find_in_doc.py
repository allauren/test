from docx import Document
from nltk.stem import PorterStemmer
from nltk.tokenize import sent_tokenize, word_tokenize
import re
import sys
from docx.enum.text import WD_COLOR_INDEX
from docx import Document
from docx.text.run import Font, Run
from docx.shared import RGBColor
import copy
import cloudconvert


def findWholeWord(w):
    return re.compile(r'\b({0})\b'.format(w), flags=re.IGNORECASE).search

def tokenizer(line):
    ps = PorterStemmer()
    if type(line) == str :
        word_list = word_tokenize(line)
    else :
        word_list = line
    token = []
    for w in word_list:
        token.append([ps.stem(w), w])
    return token

def check_words(file, words):
    words = words.split(',')
    founds = 0
    document = Document(file)
    newdoc = Document()
    twords = tokenizer(words)
    for paragraph in document.paragraphs :
        tsentence = tokenizer(paragraph.text)
        for word in tsentence :
            for fword in twords :
                if fword[0]== word[0]:
                    founds += 1
    sys.stdout.write(str(founds))

if __name__ == '__main__':
    check_words(sys.argv[1], sys.argv[2])
